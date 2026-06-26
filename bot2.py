# bot.py
import sqlite3
import random
import warnings
from datetime import datetime, timedelta
import os

from telegram import (
    Update,
    ReplyKeyboardMarkup,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
)
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    ConversationHandler,
    filters,
)

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
import asyncio

warnings.filterwarnings("ignore", category=UserWarning)

# ================== SOZLAMALAR ==================
TOKEN = "8377937512:AAG7ymiz1sQg0WlePvBqqjGxgBwrYEkumqU"
BOSS_ID = 8411652081

WORK_START = "09:00"
LATE_FINE_PER_MINUTE = 1000

# Conversation states
FULLNAME, PASSPORT, MAOSH = range(3)
KELDIM_CODE, PIN_STEP = range(2)

# ================== DATABASE ==================
conn = sqlite3.connect("ishchilar.db", check_same_thread=False)
cursor = conn.cursor()

cursor.execute("""
CREATE TABLE IF NOT EXISTS ishchilar (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    fullname TEXT,
    passport TEXT UNIQUE,
    code INTEGER UNIQUE,
    role TEXT DEFAULT 'worker',
    daily_salary INTEGER DEFAULT 100000,
    pin TEXT DEFAULT '0000',
    tg_id TEXT DEFAULT NULL,
    tax_percent INTEGER DEFAULT 12
)
""")

cursor.execute("""
CREATE TABLE IF NOT EXISTS attendance (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    passport TEXT,
    date TEXT,
    time TEXT,
    confirmed INTEGER DEFAULT 0
)
""")

cursor.execute("""
CREATE TABLE IF NOT EXISTS davomat (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    passport TEXT,
    date TEXT,
    time TEXT,
    confirmed INTEGER DEFAULT 0
)
""")

cursor.execute("""
CREATE TABLE IF NOT EXISTS moliya (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    passport TEXT,
    date TEXT,
    late_minutes INTEGER DEFAULT 0,
    jarima INTEGER DEFAULT 0,
    avans INTEGER DEFAULT 0
)
""")

cursor.execute("""
CREATE TABLE IF NOT EXISTS messages (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    ishchi_id INTEGER,
    passport TEXT,
    sender TEXT,   -- 'boss' yoki 'worker'
    type TEXT,     -- 'text' yoki 'voice'
    content TEXT,
    date TEXT,
    is_read INTEGER DEFAULT 0
)
""")
conn.commit()

# Clean old davomat (optional)
cutoff = (datetime.now() - timedelta(days=60)).strftime("%Y-%m-%d")
cursor.execute("DELETE FROM davomat WHERE date < ?", (cutoff,))
cursor.execute("DELETE FROM attendance WHERE date < ?", (cutoff,))
conn.commit()


# ================== HELPERS ==================
def generate_code():
    while True:
        c = random.randint(100000, 999999)
        cursor.execute("SELECT 1 FROM ishchilar WHERE code=?", (c,))
        if not cursor.fetchone():
            return c


def get_month_days_and_attendance(month=None):
    """
    Oylik sanalar va ishchilar ro'yxati + kelgan sanalar
    return: days, [(fullname, passport, daily_salary, kelgan_sanalar), ...]
    """
    if not month:
        month = datetime.now().strftime("%Y-%m")

    conn_local = sqlite3.connect("ishchilar.db")
    cur = conn_local.cursor()
    cur.execute("SELECT fullname, passport, daily_salary FROM ishchilar WHERE role='worker'")
    workers = cur.fetchall()

    start_date = datetime.strptime(month + "-01", "%Y-%m-%d")
    end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(days=1)
    days = [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range((end_date - start_date).days + 1)]

    result = []
    for fullname, passport, salary in workers:
        cur.execute("SELECT date FROM davomat WHERE passport=? AND confirmed=1 AND date LIKE ?", (passport, f"{month}%"))
        kelgan_sanalar = {r[0] for r in cur.fetchall()}
        result.append((fullname, passport, salary, kelgan_sanalar))

    conn_local.close()
    return days, result


def export_oylik_maosh_excel_for(passport, month=None):
    """
    Berilgan passport ga oid oylik maosh hisobotini Excel fayliga yozadi va filename qaytaradi.
    Fayl: oylik_maosh_<passport>_<month>.xlsx
    """
    days, workers = get_month_days_and_attendance(month)
    month_str = month if month else datetime.now().strftime("%Y-%m")
    ish_kunlari = 26

    wb = Workbook()
    ws = wb.active
    ws.title = "Oylik maosh"
    ws.append(["Ism Familiya", "Kelgan kun", "Kelmagan kun", "Oylik maosh", "Ayirish", "Qoldi"])

    found = False
    for fullname, p, salary, kelgan_sanalar in workers:
        if p == passport:
            found = True
            kelgan = len(kelgan_sanalar)
            kelmagan = ish_kunlari - kelgan if ish_kunlari >= kelgan else 0
            bir_kunlik = salary / ish_kunlari if ish_kunlari else 0
            ayirish = kelmagan * bir_kunlik
            qoldi = salary - ayirish
            ws.append([fullname, kelgan, kelmagan, salary, int(ayirish), int(qoldi)])
            break

    if not found:
        return None

    file_name = f"oylik_maosh_{passport}_{month_str}.xlsx"
    wb.save(file_name)
    return file_name


def export_worker_davomat_excel(passport, month=None):
    """
    Berilgan passport ga oid oylik davomad jadvalini Excel faylga yozadi
    """
    if not month:
        month = datetime.now().strftime("%Y-%m")

    start_date = datetime.strptime(month + "-01", "%Y-%m-%d")
    end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(days=1)
    days = [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range((end_date - start_date).days + 1)]

    conn_local = sqlite3.connect("ishchilar.db")
    cur = conn_local.cursor()
    cur.execute("SELECT fullname FROM ishchilar WHERE passport=?", (passport,))
    row = cur.fetchone()
    if not row:
        conn_local.close()
        return None
    fullname = row[0]

    wb = Workbook()
    ws = wb.active
    ws.title = "Davomat"
    # Header
    ws.append(["Ism Familiya"] + [d[-2:] for d in days] + ["Kelgan", "Kelmagan"])

    cur.execute("SELECT date FROM davomat WHERE passport=? AND confirmed=1 AND date LIKE ?", (passport, f"{month}%"))
    kelgan = {r[0] for r in cur.fetchall()}

    row = [fullname] + ["✅" if d in kelgan else "❌" for d in days]
    kelgan_cnt = len(kelgan)
    kelmagan_cnt = len(days) - kelgan_cnt
    row += [kelgan_cnt, kelmagan_cnt]
    ws.append(row)

    file_name = f"davomat_{passport}_{month}.xlsx"
    wb.save(file_name)
    conn_local.close()
    return file_name


# ================== START / ROLE ==================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    kb = ReplyKeyboardMarkup([["👔 Boshliq"], ["👷 Ishchi"]], resize_keyboard=True)
    await update.message.reply_text("Kim sifatida kirasiz?", reply_markup=kb)


async def role_handler(update, context):
    text = update.message.text or ""
    if "Boshliq" in text:
        if update.effective_user.id != BOSS_ID:
            await update.message.reply_text("⛔ Siz boshliq emassiz")
            return ConversationHandler.END
        context.user_data["role"] = "boshliq"   # <<< Qo‘shildi
        await boshliq_menu(update, context)
        return ConversationHandler.END
    else:
        context.user_data["role"] = "ishchi"    # <<< Qo‘shildi
        await ishchi_main_menu(update, context)
        return ConversationHandler.END


# ================== MENULAR ==================
async def boshliq_menu(update, context):
    kb = ReplyKeyboardMarkup(
        [
            ["📨 Ish berish", "💰 Oylik maosh"],
            ["📄 PDF chiqarish", "📊 Excel chiqarish"],
            ["📆 Kunlik", "📆 Oylik"],
            ["➕ Ishchi qo‘shish", "❌ Ishchi o‘chirish"],
            ["⬅️ Orqaga"]
        ],
        resize_keyboard=True
    )
    if update.callback_query:
        await update.callback_query.message.reply_text("👔 Boshliq paneli", reply_markup=kb)
    else:
        await update.message.reply_text("👔 Boshliq paneli", reply_markup=kb)


async def ishchi_main_menu(update, context):
    kb = ReplyKeyboardMarkup(
        [
            ["✅ Keldim", "💰 Oylik maosh"],
            ["📋 Davomat", "📝 Xabar"],
            ["⬅️ Orqaga"]
        ],
        resize_keyboard=True
    )
    if update.callback_query:
        await update.callback_query.message.reply_text("👷 Ishchi paneli", reply_markup=kb)
    else:
        await update.message.reply_text("👷 Ishchi paneli", reply_markup=kb)

# Ishchi panelidagi ortga
async def ishchi_orqaga(update, context):
    """Ishchi panelidagi '⬅️ Orqaga' tugmasi ishchi main menu ga qaytaradi"""
    await ishchi_main_menu(update, context)

# Boshliq panelidagi ortga
async def boshliq_orqaga(update, context):
    """Boshliq panelidagi '⬅️ Orqaga' tugmasi boshliq menu ga qaytaradi"""
    await boshliq_menu(update, context)


# ================== ISHCHI: KELDIM ==================
async def men_keldim(update, context):
    q = update.message
    today = datetime.now().strftime("%Y-%m-%d")
    now_time = datetime.now().strftime("%H:%M")

    passport = context.user_data.get("passport")
    if not passport:
        await q.reply_text("❌ Passport topilmadi. /start ni bosing va kod bilan kiring.")
        return

    conn_local = sqlite3.connect("ishchilar.db")
    cursor_local = conn_local.cursor()

    cursor_local.execute(
        "SELECT 1 FROM davomat WHERE passport=? AND date=?",
        (passport, today)
    )
    if cursor_local.fetchone():
        await q.reply_text("✅ Siz bugun kelgan deb belgilab bo‘lgansiz")
        conn_local.close()
        return

    cursor_local.execute(
        "INSERT INTO davomat (passport, date, time, confirmed) VALUES (?, ?, ?, 1)",
        (passport, today, now_time)
    )

    try:
        work_start_dt = datetime.strptime(WORK_START, "%H:%M")
        now_dt = datetime.strptime(now_time, "%H:%M")
        late_minutes = max(0, int((now_dt - work_start_dt).total_seconds() // 60))
        jarima = late_minutes * LATE_FINE_PER_MINUTE
    except Exception:
        late_minutes, jarima = 0, 0

    cursor_local.execute("""
        INSERT INTO moliya (passport, date, late_minutes, jarima)
        VALUES (?, ?, ?, ?)
    """, (passport, today, late_minutes, jarima))

    conn_local.commit()
    await q.reply_text(
        f"🟢 Kelganingiz qayd etildi!\n"
        f"🔴 Kech qolgan daqiqalar: {late_minutes}\n"
        f"💸 Jarima: {jarima} so'm"
    )
    conn_local.close()


# ================== ISHCHI: KOD + PIN ==================
async def keldim_code_start(update, context):
    kb = ReplyKeyboardMarkup([["⬅️ Orqaga"]], resize_keyboard=True)
    await update.message.reply_text(
        "🟢 Iltimos, ishchi kodingizni kiriting:",
        reply_markup=kb
    )
    return KELDIM_CODE


async def check_code_for_keldim(update, context):
    code_text = update.message.text.strip()
    if code_text == "⬅️ Orqaga":
        await ishchi_main_menu(update, context)
        return ConversationHandler.END

    try:
        code = int(code_text)
    except ValueError:
        await update.message.reply_text("❌ Iltimos, faqat raqam kiriting")
        return KELDIM_CODE

    cursor.execute("SELECT passport, fullname FROM ishchilar WHERE code=?", (code,))
    result = cursor.fetchone()
    if not result:
        await update.message.reply_text("❌ Kod topilmadi, qayta kiriting")
        return KELDIM_CODE

    passport, fullname = result
    context.user_data["passport"] = passport
    context.user_data["fullname"] = fullname

    await update.message.reply_text(f"👋 {fullname}, xush kelibsiz!\n🔒 Iltimos, PIN kodni kiriting:")
    return PIN_STEP


async def check_pin_for_keldim(update, context):
    pin = update.message.text.strip()
    passport = context.user_data.get("passport")
    cursor.execute("SELECT fullname FROM ishchilar WHERE passport=? AND pin=?", (passport, pin))
    row = cursor.fetchone()
    if not row:
        await update.message.reply_text("❌ PIN noto‘g‘ri, qayta kiriting")
        return PIN_STEP

    fullname = row[0]
    context.user_data["fullname"] = fullname

    # Keldimni saqlaymiz
    await men_keldim(update, context)

    # Ishchi panelini ko'rsatamiz
    await update.message.reply_text(f"✅ Davomat saqlandi!\n👤 {fullname}")
    await ishchi_main_menu(update, context)
    return ConversationHandler.END


conv_keldim = ConversationHandler(
    entry_points=[MessageHandler(filters.Regex("^✅ Keldim$") & ~filters.COMMAND, keldim_code_start)],
    states={
        KELDIM_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, check_code_for_keldim)],
        PIN_STEP: [MessageHandler(filters.TEXT & ~filters.COMMAND, check_pin_for_keldim)],
    },
    fallbacks=[MessageHandler(filters.Regex("^⬅️ Orqaga$"), ishchi_main_menu)]
)


# ================== ISHCHI: Oylik maosh (ISHCHI TOMONI) ==================
async def ishchi_oylik_maosh_request(update, context):
    """
    Ishchi panelidagi '💰 Oylik maosh' tugmasi -> shu user (context.user_data passport) uchun excel yaratib yuboradi.
    Agar passport user_data da bo'lmasa, iltimos kod bilan kiring deb so'raydi.
    """
    passport = context.user_data.get("passport")
    if not passport:
        await update.message.reply_text("❌ Siz hali tizimga kirmadingiz. Iltimos avval ✅ Keldim orqali kod va PIN bilan kiring.")
        return

    file = export_oylik_maosh_excel_for(passport)
    if not file:
        await update.message.reply_text("❌ Ma'lumot topilmadi.")
        return

    await update.message.reply_document(open(file, "rb"), caption="💰 Sizning oylik maosh hisobotingiz (Excel)")
    os.remove(file)


# ================== ISHCHI: DAVOMAT (ISHCHI TOMONI) ==================
async def ishchi_davomat_request(update, context):
    passport = context.user_data.get("passport")
    if not passport:
        await update.message.reply_text("❌ Siz tizimga kirmagansiz. Iltimos avval ✅ Keldim orqali kod va PIN bilan kiring.")
        return

    file = export_worker_davomat_excel(passport)
    if not file:
        await update.message.reply_text("❌ Davomat topilmadi.")
        return

    await update.message.reply_document(open(file, "rb"), caption="📋 Oylik davomat (Excel)")
    os.remove(file)


# ================== ISHCHI: XABAR (ISHCHI TOMONI) ==================
async def ishchi_xabar_menu(update, context):
    """
    Ishchi '📝 Xabar' ni bosganda bossdan kelgan xabarlarni ko'rsatish va yangi xabar yuborish tugmasini beradi.
    """
    passport = context.user_data.get("passport")
    if not passport:
        await update.message.reply_text("❌ Siz tizimga kirmagansiz. Iltimos avval ✅ Keldim orqali kod va PIN bilan kiring.")
        return

    conn_local = sqlite3.connect("ishchilar.db")
    cur = conn_local.cursor()
    # Bossdan kelgan xabarlar
    cur.execute("SELECT sender, type, content, date, is_read FROM messages WHERE passport=? AND sender='boss' ORDER BY id DESC LIMIT 10", (passport,))
    rows = cur.fetchall()
    conn_local.close()

    if not rows:
        text = "📭 Sizga boss tomonidan hali xabar yuborilmagan."
    else:
        text = "📬 Oxirgi xabarlar (bossdan):\n\n"
        for sender, t, content, date, is_read in rows:
            prefix = "✓" if is_read else "✉️"
            if t == "text":
                text += f"{prefix} {date}\n{content}\n\n"
            else:
                text += f"{prefix} {date} (voice)\n\n"

    kb = ReplyKeyboardMarkup([["✉️ Yangi xabar"], ["⬅️ Orqaga"]], resize_keyboard=True)
    await update.message.reply_text(text, reply_markup=kb)


async def ishchi_send_message_to_boss(update, context):
    """
    Ishchi tomonidan bossga yozilgan xabarni saqlaydi (messages jadvaliga) va bossga yuboradi.
    """
    passport = context.user_data.get("passport")
    fullname = context.user_data.get("fullname", "NoName")
    text = update.message.text
    if text == "✉️ Yangi xabar":
        await update.message.reply_text("📨 Xabar matnini yuboring:")
        # state handled by next message - we will store a flag
        context.user_data["awaiting_worker_msg"] = True
        return

    if context.user_data.get("awaiting_worker_msg"):
        # save message
        conn_local = sqlite3.connect("ishchilar.db")
        cur = conn_local.cursor()
        cur.execute("SELECT id FROM ishchilar WHERE passport=?", (passport,))
        r = cur.fetchone()
        ishchi_id = r[0] if r else None
        cur.execute(
            "INSERT INTO messages (ishchi_id, passport, sender, type, content, date, is_read) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (ishchi_id, passport, "worker", "text", f"{fullname}:\n{text}", datetime.now().strftime("%Y-%m-%d %H:%M"), 0)
        )
        conn_local.commit()
        conn_local.close()
        context.user_data["awaiting_worker_msg"] = False

        # send to boss (if you want)
        try:
            await context.bot.send_message(chat_id=BOSS_ID, text=f"📢 Ishchidan xabar ({fullname} / {passport}):\n\n{text}")
        except Exception:
            pass

        await update.message.reply_text("✅ Xabar boshliqqa yuborildi.")
        await ishchi_main_menu(update, context)


# ================== BOSHLIQ: ISH BERISH (MANAGING) ==================
from telegram import (
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    ReplyKeyboardMarkup
)
from telegram.ext import ConversationHandler, CallbackQueryHandler, MessageHandler, filters
from telegram.constants import ChatAction
import sqlite3

ISHBERISH_SELECT, ISHBERISH_MESSAGE = range(2)


async def ish_berish_start(update, context):
    """👔 Boshliq ish berishni boshlaydi — ishchilar ro‘yxati"""
    if update.effective_user.id != BOSS_ID:
        await update.message.reply_text("⛔ Sizda ruxsat yo‘q")
        return ConversationHandler.END

    conn = sqlite3.connect("ishchilar.db")
    cursor = conn.cursor()
    cursor.execute("SELECT id, fullname FROM ishchilar WHERE role='worker'")
    workers = cursor.fetchall()
    conn.close()

    if not workers:
        await update.message.reply_text("❌ Ishchilar topilmadi.")
        return ConversationHandler.END

    buttons = [
        [InlineKeyboardButton(name, callback_data=f"ish_{wid}")]
        for wid, name in workers
    ]
    buttons.append([InlineKeyboardButton("⬅️ Orqaga", callback_data="back_boss")])

    await update.message.reply_text(
        "📋 Ishchi tanlang:",
        reply_markup=InlineKeyboardMarkup(buttons)
    )
    return ISHBERISH_SELECT


async def ish_berish_tanlash(update, context):
    """Ishchi tanlangandan keyin xabar so‘raladi"""
    q = update.callback_query
    await q.answer()

    if q.data == "back_boss":
        await boshliq_menu(q, context)
        return ConversationHandler.END

    ishchi_id = int(q.data.split("_")[1])
    context.user_data["ishchi_id"] = ishchi_id

    conn = sqlite3.connect("ishchilar.db")
    cursor = conn.cursor()
    cursor.execute("SELECT fullname, tg_id FROM ishchilar WHERE id=?", (ishchi_id,))
    row = cursor.fetchone()
    conn.close()

    fullname, tg_id = row
    context.user_data["ishchi_tg_id"] = tg_id

    await q.message.reply_text(
        f"✍️ {fullname} uchun xabar yozing yoki 🎤 ovoz yuboring.\n\n"
        f"Tugatgach, “📤 Yuborish” ni bosing.",
        reply_markup=ReplyKeyboardMarkup(
            [["📤 Yuborish"], ["⬅️ Bekor qilish"]],
            resize_keyboard=True
        )
    )
    return ISHBERISH_MESSAGE


async def ish_berish_yuborish(update, context):
    """Xabarni ishchiga yuborish"""
    tg_id = context.user_data.get("ishchi_tg_id")
    if not tg_id:
        await update.message.reply_text("❌ Ishchi tanlanmagan.")
        return ConversationHandler.END

    await context.bot.send_chat_action(tg_id, ChatAction.TYPING)

    if update.message.voice:
        await context.bot.send_voice(
            chat_id=tg_id,
            voice=update.message.voice.file_id,
            caption="📢 Boshliqdan ovozli xabar"
        )
    else:
        text = update.message.text
        await context.bot.send_message(
            chat_id=tg_id,
            text=f"📢 Boshliqdan xabar:\n\n{text}"
        )

    await update.message.reply_text("✅ Xabar ishchiga yuborildi.")
    await boshliq_menu(update, context)
    return ConversationHandler.END


async def ish_berish_cancel(update, context):
    await boshliq_menu(update, context)
    return ConversationHandler.END

# ================== PDF/EXCEL/WORD REPORT FUNCTIONS (boss) ==================
def word_month_report(month=None):
    if not month:
        month = datetime.now().strftime("%Y-%m")

    doc = Document()
    doc.add_heading(f"{month} oylik davomad hisobot", level=1)
    conn_local = sqlite3.connect("ishchilar.db")
    cursor_local = conn_local.cursor()
    cursor_local.execute("SELECT fullname, passport FROM ishchilar WHERE role='worker'")
    workers = cursor_local.fetchall()

    start_date = datetime.strptime(month + "-01", "%Y-%m-%d")
    end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(days=1)
    days = [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range((end_date - start_date).days + 1)]

    table = doc.add_table(rows=1, cols=2 + len(days))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ism'
    hdr_cells[1].text = 'Familiya'
    for i, d in enumerate(days):
        hdr_cells[i+2].text = d[-2:]

    for n, p in workers:
        parts = n.split(" ")
        ism = parts[0]
        familya = parts[1] if len(parts) > 1 else ""
        row_cells = table.add_row().cells
        row_cells[0].text = ism
        row_cells[1].text = familya

        kelgan_sanalar = set()
        cursor_local.execute("SELECT date FROM davomat WHERE passport=? AND date LIKE ?", (p, f"{month}%"))
        for r in cursor_local.fetchall():
            kelgan_sanalar.add(r[0])

        for i, d in enumerate(days):
            row_cells[i+2].text = "✅" if d in kelgan_sanalar else "❌"

        total_kelgan = len(kelgan_sanalar)
        total_kelmadi = len(days) - total_kelgan
        # Qo‘shimcha oxirgi ustun
        row_cells.append().text = f"{total_kelgan} kelgan, {total_kelmadi} kelmagan"

    conn_local.close()
    file = f"hisobot_{month}.docx"
    doc.save(file)
    return file


def pdf_month_report(month=None):
    days, workers = get_month_days_and_attendance(month)
    month_str = month if month else datetime.now().strftime("%Y-%m")
    file_name = f"hisobot_{month_str}.pdf"

    c = canvas.Canvas(file_name, pagesize=A4)
    y = 800

    for fullname, _, _, kelgan_sanalar in workers:
        line = f"{fullname}: " + "".join(["✅" if d in kelgan_sanalar else "❌" for d in days])
        total = f" (Kelgan: {len(kelgan_sanalar)}, Kelmagan: {len(days) - len(kelgan_sanalar)})"
        c.drawString(50, y, line + total)
        y -= 15
        if y < 50:
            c.showPage()
            y = 800

    c.save()
    return file_name


def export_oylik_maosh_excel(month=None):
    days, workers = get_month_days_and_attendance(month)
    month_str = month if month else datetime.now().strftime("%Y-%m")
    ish_kunlari = 26  # oylik ish kunlari

    wb = Workbook()
    ws = wb.active
    ws.title = "Oylik maosh"
    ws.append(["Ism Familiya", "Kelgan kun", "Kelmagan kun", "Oylik maosh", "Ayirish", "Qoldi"])

    for fullname, passport, salary, kelgan_sanalar in workers:
        kelgan = len(kelgan_sanalar)
        kelmagan = ish_kunlari - kelgan
        bir_kunlik = salary / ish_kunlari
        ayirish = kelmagan * bir_kunlik
        qoldi = salary - ayirish

        ws.append([fullname, kelgan, kelmagan, salary, int(ayirish), int(qoldi)])

    file_name = f"oylik_maosh_{month_str}.xlsx"
    wb.save(file_name)
    return file_name


# ================== ASYNC HANDLERLAR TO SEND FILES ==================
async def send_word_report(update, context):
    file = word_month_report()
    await update.message.reply_document(open(file, "rb"))
    os.remove(file)


async def send_pdf_report(update, context):
    file = pdf_month_report()
    await update.message.reply_document(open(file, "rb"))
    os.remove(file)


async def send_oylik_maosh_excel(update, context):
    if update.effective_user.id != BOSS_ID:
        await update.message.reply_text("⛔ Sizda ruxsat yo‘q")
        return

    file = export_oylik_maosh_excel()
    await update.message.reply_document(
        document=open(file, "rb"),
        caption="💰 Oylik maosh hisoboti (Excel)"
    )
    os.remove(file)


# ================== ISHCHI QO‘SHISH / O‘CHIRISH ==================
async def ishchi_qoshish(update, context):
    """Ishchi qo‘shish jarayonini boshlaydi"""
    await update.message.reply_text("Ism va familiyasini kiriting:")
    return FULLNAME


async def ishchi_fullname(update, context):
    """Ism va familiya qabul qilinadi"""
    context.user_data["fullname"] = update.message.text.title()
    await update.message.reply_text("Pasport seriya va raqamini kiriting:")
    return PASSPORT


async def ishchi_passport(update, context):
    """Pasport raqami qabul qilinadi"""
    context.user_data["passport"] = update.message.text.upper()
    await update.message.reply_text("Oylik maoshini kiriting (raqam):")
    return MAOSH


async def ishchi_maosh(update, context):
    """Oylik maosh qabul qilinadi va ishchi bazaga qo‘shiladi"""
    fullname = context.user_data["fullname"]
    passport = context.user_data["passport"]

    try:
        salary = int(update.message.text)
    except ValueError:
        await update.message.reply_text("❌ Iltimos faqat raqam kiriting")
        return MAOSH

    # ✅ Pasport takrorlanmasligini tekshiramiz
    cursor.execute("SELECT id FROM ishchilar WHERE passport=?", (passport,))
    if cursor.fetchone():
        await update.message.reply_text("⚠ Bu pasport raqami allaqachon ro‘yxatda mavjud!")
        return PASSPORT

    code = generate_code()
    cursor.execute(
        "INSERT INTO ishchilar (fullname, passport, code, role, daily_salary) VALUES (?, ?, ?, 'worker', ?)",
        (fullname, passport, code, salary)
    )
    conn.commit()

    await update.message.reply_text(
        f"✅ Ishchi qo‘shildi!\n"
        f"👤 {fullname}\n"
        f"🪪 {passport}\n"
        f"💰 {salary} so‘m\n"
        f"🔐 Kod: {code}"
    )
    return ConversationHandler.END


# === Conversation (ishchi qo‘shish) ===
conv_add_worker = ConversationHandler(
    entry_points=[MessageHandler(filters.Regex("^➕ Ishchi qo‘shish$"), ishchi_qoshish)],
    states={
        FULLNAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, ishchi_fullname)],
        PASSPORT: [MessageHandler(filters.TEXT & ~filters.COMMAND, ishchi_passport)],
        MAOSH: [MessageHandler(filters.TEXT & ~filters.COMMAND, ishchi_maosh)],
    },
    fallbacks=[MessageHandler(filters.Regex("^⬅️ Orqaga$"), boshliq_menu)]
)


# === Ishchi o‘chirish ===
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
import asyncio

# === Ishchi o‘chirish menyusi ===
async def ishchi_ochirish(update, context):
    """Ishchilar ro‘yxatini Inline tugmalar bilan ko‘rsatadi va o‘chirish imkonini beradi"""
    cursor.execute("SELECT id, fullname FROM ishchilar WHERE role='worker'")
    workers = cursor.fetchall()

    if not workers:
        await update.message.reply_text("Ishchilar mavjud emas ❌")
        return

    buttons = [
        [InlineKeyboardButton(f"{n}", callback_data=f"del_{i}")]
        for i, n in workers
    ]
    buttons.append([InlineKeyboardButton("⬅️ Orqaga", callback_data="back_boss")])

    # Callback bo‘lsa eski xabarni edit qilamiz, aks holda yangi xabar
    if update.callback_query:
        await update.callback_query.message.edit_text(
            "🗑 O‘chirish uchun ishchini tanlang:",
            reply_markup=InlineKeyboardMarkup(buttons)
        )
    else:
        await update.message.reply_text(
            "🗑 O‘chirish uchun ishchini tanlang:",
            reply_markup=InlineKeyboardMarkup(buttons)
        )


# === Ishchini o‘chirish amali ===
async def delete_worker(update, context):
    """Tanlangan ishchini bazadan o‘chiradi va Inline menyuni yangilaydi"""
    q = update.callback_query
    await q.answer()

    # 🔙 Orqaga bosilsa
    if q.data == "back_boss":
        await boshliq_menu(update, context)
        return

    # ID ni ajratib olish
    try:
        worker_id = int(q.data.split("_")[1])
    except (IndexError, ValueError):
        await q.message.reply_text("❌ Xato ma'lumot.")
        return

    # 🔥 Ishchi ma'lumotini olish
    cursor.execute("SELECT fullname FROM ishchilar WHERE id=?", (worker_id,))
    result = cursor.fetchone()
    if not result:
        await q.message.reply_text("❌ Ishchi topilmadi.")
        return
    fullname = result[0]

    # 🔥 Bazadan o‘chirish
    cursor.execute("DELETE FROM ishchilar WHERE id=?", (worker_id,))
    conn.commit()

    # ✅ Xabar chiqarish
    await q.answer(text=f"✅ Ishchi {fullname} o‘chirildi!", show_alert=True)

    # 🔁 Inline menyuni yangilash, qolgan ishchilarni ko‘rsatish
    await ishchi_ochirish(update, context)


# ================== DAVOMAD (BOSHLQ PANELI) ==================
async def davomad(update, context):
    q = update.callback_query if update.callback_query else update.message
    today = datetime.now().strftime("%Y-%m-%d")
    now = datetime.now().strftime("%H:%M")

    conn_local = sqlite3.connect("ishchilar.db")
    cursor_local = conn_local.cursor()

    cursor_local.execute("SELECT id, fullname, passport FROM ishchilar WHERE role='worker'")
    workers = cursor_local.fetchall()
    if not workers:
        await q.reply_text("Ishchilar mavjud emas ❌")
        conn_local.close()
        return

    buttons = []
    for _, name, passport in workers:
        cursor_local.execute("SELECT 1 FROM davomat WHERE passport=? AND date=?", (passport, today))
        marked = cursor_local.fetchone()
        mark = "✅" if marked else "⬜"
        buttons.append([InlineKeyboardButton(f"{mark} {name}", callback_data=f"mark_{passport.strip()}")])

    buttons.append([InlineKeyboardButton("💾 Saqlash", callback_data="save")])
    buttons.append([InlineKeyboardButton("⬅️ Orqaga", callback_data="back_boss")])

    if update.callback_query:
        await update.callback_query.edit_message_text(
            f"📋 Bugungi davomad ({today})",
            reply_markup=InlineKeyboardMarkup(buttons)
        )
    else:
        await q.reply_text(
            f"📋 Bugungi davomad ({today})",
            reply_markup=InlineKeyboardMarkup(buttons)
        )
    conn_local.close()


async def mark(update, context):
    q = update.callback_query
    await q.answer()
    today = datetime.now().strftime("%Y-%m-%d")
    now = datetime.now().strftime("%H:%M")

    conn_local = sqlite3.connect("ishchilar.db")
    cursor_local = conn_local.cursor()

    if q.data == "back_boss":
        await boshliq_menu(q, context)
        conn_local.close()
        return

    if q.data == "save":
        cursor_local.execute("UPDATE davomat SET confirmed=1 WHERE date=? AND confirmed=0", (today,))
        conn_local.commit()
        await q.message.reply_text("✅ Davomad saqlandi! (60 kun davomida saqlanadi)")
        conn_local.close()
        return

    passport = q.data.split("_")[1].strip()
    cursor_local.execute("SELECT 1 FROM davomat WHERE passport=? AND date=?", (passport, today))
    if cursor_local.fetchone():
        cursor_local.execute("DELETE FROM davomat WHERE passport=? AND date=?", (passport, today))
    else:
        cursor_local.execute(
            "INSERT INTO davomat (passport, date, time, confirmed) VALUES (?, ?, ?, 0)",
            (passport, today, now)
        )
    conn_local.commit()
    conn_local.close()
    await davomad(update, context)

# ================== KUNLIK (BOSHLQ PANELI) ==================
async def kunlik(update, context):
    today = datetime.now().strftime("%Y-%m-%d")

    conn_local = sqlite3.connect("ishchilar.db")
    cursor_local = conn_local.cursor()

    # Ishchilar va bugungi davomatni olish
    cursor_local.execute("SELECT fullname, passport FROM ishchilar WHERE role='worker'")
    all_workers = cursor_local.fetchall()

    cursor_local.execute("SELECT passport FROM davomat WHERE date=? AND confirmed=1", (today,))
    present = [r[0] for r in cursor_local.fetchall()]
    conn_local.close()

    # Matnni shakllantirish
    text = f"🗓 <b>Bugungi davomad — {today}</b>\n\n"

    kelganlar = [n for n, p in all_workers if p in present]
    kelmaganlar = [n for n, p in all_workers if p not in present]

    if kelganlar:
        text += "🟢 <b>KELGANLAR:</b>\n" + "\n".join(f"✅ {n}" for n in kelganlar) + "\n\n"
    else:
        text += "🟢 <b>KELGANLAR:</b> — yo‘q\n\n"

    if kelmaganlar:
        text += "🔴 <b>KELMAGANLAR:</b>\n" + "\n".join(f"❌ {n}" for n in kelmaganlar)
    else:
        text += "🔴 <b>KELMAGANLAR:</b> — yo‘q"

    # Natijani chiqarish
    await update.message.reply_text(text, parse_mode="HTML")

# ================== EXCEL OYLIK HISOBOT ==================
def update_excel_report(month=None):
    """Oylik davomadni Excel faylga saqlaydi va fayl nomini qaytaradi"""
    if not month:
        month = datetime.now().strftime("%Y-%m")

    file = f"hisobot_{month}.xlsx"
    if os.path.exists(file):
        wb = load_workbook(file)
    else:
        wb = Workbook()
    ws = wb.active
    ws.title = "Davomad"

    conn = sqlite3.connect("ishchilar.db")
    cursor = conn.cursor()
    cursor.execute("SELECT fullname, passport FROM ishchilar WHERE role='worker'")
    workers = cursor.fetchall()

    start_date = datetime.strptime(month + "-01", "%Y-%m-%d")
    end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(days=1)
    days = [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range((end_date - start_date).days + 1)]

    ws.append(["Ism", "Familiya"] + [d[-2:] for d in days] + ["Kelgan", "Kelmagan"])

    for n, p in workers:
        parts = n.split(" ")
        ism = parts[0]
        familya = parts[1] if len(parts) > 1 else ""
        kelgan_sanalar = set()
        cursor.execute("SELECT date FROM davomat WHERE passport=? AND date LIKE ?", (p, f"{month}%"))
        for r in cursor.fetchall():
            kelgan_sanalar.add(r[0])
        row = [ism, familya] + ["✅" if d in kelgan_sanalar else "❌" for d in days]
        row += [len(kelgan_sanalar), len(days) - len(kelgan_sanalar)]
        ws.append(row)

    conn.close()
    wb.save(file)
    return file


# ================== OYLIK DAVOMAD (EXCEL) ==================
async def oylik_davomad(update, context):
    """Oylik davomadni Excel ko‘rinishida chiqaradi"""
    if update.effective_user.id != BOSS_ID:
        await update.message.reply_text("⛔ Sizda ruxsat yo‘q")
        return

    file_name = update_excel_report()  # Excel fayl yaratish
    await update.message.reply_document(open(file_name, "rb"), caption="📊 Oylik davomad hisobot (Excel)")
    os.remove(file_name)


import os
from datetime import datetime, timedelta
from openpyxl import Workbook, load_workbook
import sqlite3

# ================== EXCEL OYLIK HISOBOT ==================
def update_excel_report(month=None):
    """Oylik davomadni Excel faylga saqlaydi va fayl nomini qaytaradi"""
    if not month:
        month = datetime.now().strftime("%Y-%m")

    file = f"hisobot_{month}.xlsx"
    if os.path.exists(file):
        wb = load_workbook(file)
    else:
        wb = Workbook()
    ws = wb.active
    ws.title = "Davomad"

    conn = sqlite3.connect("ishchilar.db")
    cursor = conn.cursor()
    cursor.execute("SELECT fullname, passport FROM ishchilar WHERE role='worker'")
    workers = cursor.fetchall()

    # Sana ustunlarini yaratish
    start_date = datetime.strptime(month + "-01", "%Y-%m-%d")
    end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(days=1)
    days = [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range((end_date - start_date).days + 1)]

    ws.append(["Ism", "Familiya"] + [d[-2:] for d in days] + ["Kelgan", "Kelmagan"])

    for n, p in workers:
        parts = n.split(" ")
        ism = parts[0]
        familya = parts[1] if len(parts) > 1 else ""

        kelgan_sanalar = set()
        cursor.execute("SELECT date FROM davomat WHERE passport=? AND date LIKE ?", (p, f"{month}%"))
        for r in cursor.fetchall():
            kelgan_sanalar.add(r[0])

        row = [ism, familya] + ["✅" if d in kelgan_sanalar else "❌" for d in days]
        row += [len(kelgan_sanalar), len(days) - len(kelgan_sanalar)]
        ws.append(row)

    conn.close()
    wb.save(file)
    return file

# ================== OYLIK DAVOMAD (EXCEL) ==================
async def oylik_davomad(update, context):
    """Oylik davomadni Excel ko‘rinishida chiqaradi"""
    if update.effective_user.id != BOSS_ID:
        await update.message.reply_text("⛔ Sizda ruxsat yo‘q")
        return

    # Excel faylni yaratish
    file_name = update_excel_report()  # Endi aniqlangan funksiya ishlaydi

    # Telegramga jo'natish
    await update.message.reply_document(
        open(file_name, "rb"),
        caption="📊 Oylik davomad hisobot (Excel)"
    )

    # Faylni tizimdan o‘chirish
    os.remove(file_name)


# ================== HANDLERLAR / MAIN ==================
def main():
    app = ApplicationBuilder().token(TOKEN).build()

    # ====== Start / role tanlash ======
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Regex("^(👔 Boshliq|👷 Ishchi)$"), role_handler))

    # ====== Boshliq: Ish berish conversation ======
    conv_ish_berish = ConversationHandler(
        entry_points=[MessageHandler(filters.Regex("^📨 Ish berish$"), ish_berish_start)],
        states={
            ISHBERISH_SELECT: [CallbackQueryHandler(ish_berish_tanlash, pattern="^ish_")],
            ISHBERISH_MESSAGE: [
                MessageHandler(filters.TEXT | filters.VOICE, ish_berish_yuborish),
            ],
        },
        fallbacks=[MessageHandler(filters.Regex("^⬅️ Bekor qilish$"), ish_berish_cancel)],
    )
    app.add_handler(conv_ish_berish)
        # Ishchi xabar yuborish (faqat COMMAND bo‘lmagan xabarlar)
    app.add_handler(
        MessageHandler(
            filters.Regex("^📝 Xabar$") & ~filters.COMMAND,
            ishchi_xabar_yuborish
        )
    )
    # ====== Ishchi qo‘shish conversation ======
    app.add_handler(conv_add_worker)

    # ====== Ishchi "Keldim" conversation ======
    app.add_handler(conv_keldim)

    # ====== Boshliq hisobotlari ======
    app.add_handler(MessageHandler(filters.Regex("^📄 PDF chiqarish$"), send_pdf_report))
    app.add_handler(MessageHandler(filters.Regex("^💰 Oylik maosh$"), send_oylik_maosh_excel))
    app.add_handler(MessageHandler(filters.Regex("^📆 Kunlik$"), kunlik))
    app.add_handler(MessageHandler(filters.Regex("^📆 Oylik$"), oylik_davomad))

    # ====== Ishchi paneli ======
    app.add_handler(MessageHandler(filters.Regex("^💰 Oylik maosh$"), ishchi_oylik_maosh_request))
    app.add_handler(MessageHandler(filters.Regex("^📋 Davomat$"), ishchi_davomat_request))
    app.add_handler(MessageHandler(filters.Regex("^📝 Xabar$"), ishchi_xabar_menu))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, ishchi_send_message_to_boss))

    # ====== Qo‘shimcha (boshliq) ======
    app.add_handler(MessageHandler(filters.Regex("^➕ Ishchi qo‘shish$"), ishchi_qoshish))
    app.add_handler(MessageHandler(filters.Regex("^❌ Ishchi o‘chirish$"), ishchi_ochirish))

    # ====== Callback tugmalar ======
    app.add_handler(CallbackQueryHandler(delete_worker, pattern="^(del_|back_boss)$"))
    app.add_handler(CallbackQueryHandler(mark, pattern="^(mark_|save|back_boss)$"))
    app.add_handler(CallbackQueryHandler(ish_berish_tanlash, pattern="^ish_"))

    # ====== Ishchi va Boshliq "⬅️ Orqaga" tugmasi ======
    async def orqaga_handler(update, context):
        user_role = context.user_data.get("role")
        if user_role == "ishchi":
            await ishchi_main_menu(update, context)
        elif user_role == "boshliq":
            await boshliq_menu(update, context)
        else:
            await update.message.reply_text("Panel aniqlanmadi. /start ni bosing")

    app.add_handler(MessageHandler(filters.Regex("^⬅️ Orqaga$"), orqaga_handler))

    print("🤖 Bot ishga tushdi")
    app.run_polling()


if __name__ == "__main__":
    main()
